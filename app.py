from flask import Flask, request
from dotenv import load_dotenv
from contextlib import closing
from PIL import Image
from docx import Document
from docx.shared import Cm
from inline_keyboards import (
    file_formats,
)
from update_handlers import (
    get_chat_id,
    get_callback_data,
    get_message_text
)
from filters import (
    message_text_filter,
    callback_data_filter,
    is_message_filter
)
from bot import Bot
from database import Database
import requests, os, psycopg2, shutil


load_dotenv()

app = Flask(__name__) 
bot = Bot(os.getenv("BOT_TOKEN")) 
conn = psycopg2.connect(dbname="userimage", user="timur", password="1qaz2wsx1AZ", host="localhost")

@message_text_filter(["/start"])
def send_format_list():
    with closing(conn.cursor()) as cursor:
        cursor.execute(f"SELECT * FROM image WHERE UserId={get_chat_id(request)}")
        if not cursor.fetchall():
            cursor.execute(f"INSERT INTO image (UserId) VALUES ({get_chat_id(request)})")
        else:
            cursor.execute("SELECT * FROM image")
            print(cursor.fetchall())

    bot.send_photo(get_chat_id(request), "https://vk.com/photo-136259311_457346999", caption="send me a link to the photo and select the file extension you want to save it in 📑", reply_markup=file_formats)

@callback_data_filter(["jpg", "png", "ico", "jpeg", "bmp", "pdf", "docx"])
def set_file_extension_to_database():
    with closing(conn.cursor()) as cursor:
        cursor.execute(f"UPDATE image SET FileType='{get_callback_data(request)}' WHERE UserId={get_chat_id(request)}")

    bot.send_photo(get_chat_id(request), "https://vk.com/photo-197495754_457278028", caption="send a link to the photo, I'm waiting...")

@is_message_filter
def set_image_link_to_database():
    with closing(conn.cursor()) as cursor:
        cursor.execute(f"SELECT FileType FROM image WHERE UserId={get_chat_id(request)}")
        if not cursor.fetchall()[0][0]:
            return
        cursor.execute(f"UPDATE image SET Url='{get_message_text(request)}' WHERE UserId={get_chat_id(request)}") 
        cursor.execute(f"SELECT FileType FROM image WHERE UserId={get_chat_id(request)}")

        filetype = cursor.fetchone()[0]
        response = requests.get(get_message_text(request), stream=True)

        match filetype: 
            case ("png" | "jpg" | "ico" | "jpeg" | "bmp"):
                with open("images/image." + filetype, "wb") as file:
                    response.raw.decode_content = True
                    shutil.copyfileobj(response.raw, file)
            case "pdf":
                with open("images/image.png", "wb") as file:
                    response.raw.decode_content = True
                    shutil.copyfileobj(response.raw, file)

                image = Image.open("images/image.png")
                colored_image = image.convert("RGB")
                colored_image.save("images/image.pdf")
                os.remove("images/image.png")

            case "docx":
                width, height = None, None

                with open("images/image.png", "wb") as file:
                    response.raw.decode_content = True
                    shutil.copyfileobj(response.raw, file)
                    
                with Image.open("images/image.png") as file:
                    width, height = file.size
                    print(file.size)

                document = Document()
                section = document.sections[0]
                section.page_width = Cm(width)
                section.page_height = Cm(height)
                section.top_margin = Cm(0)
                section.bottom_margin = Cm(0)
                section.left_margin = Cm(0)
                section.right_margin = Cm(0)
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture("images/image.png", width=Cm(width), height=Cm(height))
                document.save("images/image.docx")
                os.remove("images/image.png")

        with open("images/image." + filetype, "rb") as file:
            bot.send_document(chat_id=get_chat_id(request), caption="come again buddy!\n/start", document=file)
        os.remove("images/image." + filetype)
        cursor.execute(f"DELETE FROM image WHERE UserId={get_chat_id(request)}")

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        print(request.json)
        send_format_list() 
        set_file_extension_to_database()
        set_image_link_to_database()

    return {"ok": True}

if __name__ == "__main__":
    app.run()

